"""生成论文 Word 版（.docx）
基于 docs/paper_draft.md 内容，用 python-docx 生成完整排版的 Word 文档。
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell):
    """给单元格加边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_font(run, name='宋体', size=10.5, bold=False, en_name='Times New Roman'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_para(doc, text, size=10.5, bold=False, align=None, indent=True, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 14, 2: 12, 3: 11}
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 11), bold=True)
    return p


def add_table(doc, headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after = Pt(4)
        run = cp.add_run(caption)
        set_font(run, size=10, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9.5, bold=True)
        set_cell_border(cell)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=9.5)
            set_cell_border(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    doc = Document()
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run('图约束推理链与子答案双向交叉核验：面向多跳推理的大模型方法')
    set_font(run, size=16, bold=True)

    # 作者
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(12)
    run = author.add_run('Anonymous Authors')
    set_font(run, size=11)

    # 摘要（英文，会议要求）
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Abstract: ')
    set_font(run, size=10.5, bold=True)
    run = p.add_run('Large language models (LLMs) suffer from error accumulation on multi-hop reasoning tasks. Some graph-enhanced variants use structural graph-validity metrics as consistency checks, but such metrics may not reflect reasoning correctness. We propose GERS, which models reasoning as a dependency DAG executed in topological order. The core contribution is bidirectional sub-answer cross-checking: using the final answer + context as an anchor, we re-derive each sub-question in reverse and compare forward/backward sub-answers, upgrading the Consistency Score from structural validity to content self-consistency. On 500 HotpotQA samples, this raises the score\'s correct/wrong discrimination from -0.0035 to +0.0847 and AUROC from 0.498 to about 0.58-0.59; the highest point-estimate configuration GERS-CV2 achieves EM=0.302, F1=0.413, with a +4pt F1 difference over CoT-SC that is significant under both a paired McNemar test on EM correctness (p=0.029) and paired bootstrap CIs excluding zero (EM [+0.006,+0.074], F1 [+0.006,+0.075]); GERS-CV2 also significantly outperforms a MoDeGraph-style graph-prompt baseline (p=0.010). We explicitly show that self-consistency is not correctness: many high-CS answers remain wrong, and 2WikiMultiHopQA exposes a boundary on deep bridge-comparison questions.')
    set_font(run, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('Keywords: ')
    set_font(run, size=10.5, bold=True)
    run = p.add_run('Large Language Models; Multi-hop Reasoning; Chain-of-Thought; Graph Representation; Bidirectional Cross-Checking; Consistency Diagnostics')
    set_font(run, size=10)

    # 1 引言
    add_heading(doc, '1 引言', 1)
    add_para(doc, '多跳推理要求模型跨多个证据片段组合推导出答案，是衡量 LLM 复杂推理能力的核心任务之一。Chain-of-Thought（CoT）提示通过引导模型输出中间推理步骤，显著提升了 LLM 的表现。然而，标准 CoT 以线性文本串接推理步骤，存在两类核心不足：')
    add_para(doc, '1. 非线性依赖缺失：CoT 以线性文本串接推理步骤，无法表达子问题间的分支、合流与并行依赖。当推理路径本应是 DAG 而非链时，线性化导致结构信息丢失。')
    add_para(doc, '2. 质量无差别投票：CoT-SC 通过多次采样取众数答案，但投票是"等权"的，无法利用推理路径本身的质量信息。')
    add_para(doc, '近期工作尝试用图结构增强推理：GoT 提出思维图支持合并与蒸馏，但不采用本文的子问题拓扑执行；RwG 将上下文结构化为实体关系图；MoDeGraph 构建多跳依赖图，属于 prompt 方法。GERS 关注当图结构参与子问题执行与校验时，图派生一致性分数是否能作为诊断与弱候选质量信号。')
    add_para(doc, '针对上述问题，本文提出 GERS，将推理显式建模为推理依赖图（DAG）并按拓扑序执行。我们发现，单纯用图论结构性质计算 Consistency Score 是无效的——LLM 生成的分解图几乎都是合法 DAG，导致所有样本得分趋同（聚集在 0.6~0.7），多路选优退化为随机抽取（500 条实测对错区分度仅 -0.0035）。为此，本文提出核心创新子答案双向交叉核验：以最终答案 + 上下文为锚，反向逐个重答每个子问题，对比正向与反向子答案的一致性，将 Consistency Score 从"图是否合法"升级为"推理内容是否自洽"。显式 DAG 分解提供了命名子问题单元和依赖边，使该核验比应用于非结构化 CoT 文本更直接、更可解释。')
    add_para(doc, '本文的主要贡献：（1）子答案双向交叉核验：将 CS 对错区分度从 -0.0035 修复到 +0.0847，AUROC 从 0.498 提升到约 0.58-0.59，使 CS 成为弱但有用的内容自洽诊断信号；（2）性能验证：点估计最高配置 GERS-CV2 在 HotpotQA 500 条上 EM=0.302、F1=0.413，F1 相比 CoT-SC 高 4pt，且在配对 McNemar 检验（p=0.029）与配对 bootstrap CI（不跨零）下均显著，并对 MoDeGraph-style 基线显著（p=0.010）；（3）诚实的适用边界：旧结构分图级多路选优无额外增益，bridge_comparison 存在错误传播局限；（4）工程贡献：指标修复、答案类型回扣、提取公平性。')

    # 2 相关工作
    add_heading(doc, '2 相关工作', 1)
    add_para(doc, '思维链推理。Wei 等提出 CoT；Wang 等提出 Self-Consistency（CoT-SC），等权投票未利用推理路径质量；Yao 等提出 Tree of Thoughts（ToT），子问题间无显式依赖建模。这些方法均基于线性或树形结构。')
    add_para(doc, '图结构增强推理。GoT 将推理建模为图但不采用本文的子问题拓扑执行；RwG 将上下文结构化为实体关系图；MoDeGraph 本质仍是 prompt 方法。GERS 将推理依赖图与执行流程耦合，并研究图派生一致性作为诊断与弱排序信号的作用。')
    add_para(doc, '一致性校验与自评估。Process Reward Models 依赖大量标注；Self-Refine 作用于线性文本、缺少可核验的子问题结构。GERS 用图论算法做结构化校验，并设计双向交叉核验，使一致性分数可被分析并用于受控变体。')
    add_table(doc, ['方法族', '结构', '执行顺序', '可核验子答案'], [
        ['CoT', '线性文本', '隐式', '否'],
        ['CoT-SC', '多条线性样本', '隐式', '否'],
        ['ToT / GoT', '树或图状thought', '搜索/灵活', '状态或thought节点'],
        ['RwG / 图提示', '实体/上下文图', 'prompt条件化', '通常否'],
        ['MoDeGraph-style', '依赖图prompt', 'prompt条件化', '实体/关系级'],
        ['GERS', '子问题DAG', '拓扑执行', '正反向核验'],
    ], caption='表 0  GERS 与代表性推理范式的定位对比')

    # 3 方法
    add_heading(doc, '3 方法', 1)
    add_heading(doc, '3.1 总体框架', 2)
    add_para(doc, 'GERS 推理流程（图1）：输入问题 Q + 上下文 C → ① 分解为子问题 {q_i} 及依赖 → ② 构建 DAG G=(V,E) → ③ 拓扑排序 τ → ④ 逐步执行生成子答案 a_i → ⑤ 汇总生成最终答案 A（答案类型回扣）→ ⑥ 计算交叉验证一致性分 S = 0.3·S_struct + 0.7·S_crossval → 输出 A + G + S。另设历史诊断基线 GERS-SC：采样 K 条 DAG 并用原始结构 CS 选优。')

    add_heading(doc, '3.2 推理状态图表示', 2)
    add_para(doc, '节点类型：FactNode（事实/原始问题）、StepNode（子问题及答案）、ConclusionNode（最终结论）。边类型：DeriveEdge（推导）、SupportEdge（支撑）、ConflictEdge（互斥）。LLM 将问题分解为子问题列表并标注依赖，据此创建节点与边。')

    add_heading(doc, '3.3 图约束链路生成', 2)
    add_para(doc, '路径规划：拓扑排序得执行顺序 τ，确保每个子问题被回答时其前驱已获得答案。逐步执行与汇总：按拓扑序逐个回答子问题，依赖传递，最后汇总生成最终答案 A。答案类型回扣：约束最终答案匹配原问题要求的实体类型（如问 film 答电影名而非中间实体导演名）。')

    add_heading(doc, '3.4 一致性校验与子答案双向交叉核验', 2)
    add_para(doc, '结构层 Consistency Score（基线）：S_struct = 0.35·C_conn + 0.30·C_cycle + 0.35·C_cov（连通性、无环性、覆盖度），含推理链长度惩罚。')
    add_para(doc, '结构层的局限：LLM 生成的分解图几乎都是合法 DAG，导致 S_struct 高度聚集（500 条均值 0.662，stdev 0.098，区分度 -0.0035），无法区分推理好坏。')
    add_para(doc, '子答案双向交叉核验（核心创新）：以"最终答案 A + 上下文 C"为锚反向逐个重答子问题，得反向子答案 {a\'_i}，对比正反向一致性：S_crossval = Σ w_i·1[match(a_i, a\'_i)] / Σ w_i。反向 Prompt 要求"基于上下文独立重答，勿照抄最终答案"以缓解 confirmation bias。显式 DAG 分解使该机制更直接、更可解释。')
    add_para(doc, 'confirmation bias 控制：由于反向验证默认可见最终答案 A，我们额外评测仅使用上下文、不使用答案锚点的反向验证变体。该变体得到 EM 0.300 / F1 0.414，CS 降至 0.646；默认设置为 EM 0.302 / F1 0.413 / CS 0.777。端任务结果基本不变，说明收益并非简单来自把最终答案喂回模型；但 CS 降低也说明答案锚点会提高内部一致性，不能把 CS 解释为事实正确性概率。')
    add_para(doc, '新 Consistency Score：S = 0.3·S_struct + 0.7·S_crossval。结构层权重降至 0.3，内容层 0.7。')

    add_heading(doc, '3.5 子答案置信度加权汇总', 2)
    add_para(doc, '每个子答案置信度由轻量启发式估计（零额外 LLM 调用），综合考量答案非空、不含不确定信号、核心实体在上下文落地、是否纯数字。低于阈值 0.5 的子答案在汇总 Prompt 中标注 [LOW CONFIDENCE]，降低单步错误传播。')

    add_heading(doc, '3.6 图级自一致性（GERS-SC）', 2)
    add_para(doc, '生成 K 条不同推理 DAG（分解温度 T∈{0.3,0.5,0.7}），用 CS 选得分最高者：A* = argmax_k S(G_k)。主表中的 GERS-SC 使用原始结构分选优；其失效正是双向交叉核验的动机。是否用新的 cross-validated score 扩展完整 K 路选优，由于会进一步放大验证成本，留作未来工作。')

    # 4 实验
    add_heading(doc, '4 实验', 1)
    add_heading(doc, '4.1 实验设置', 2)
    add_para(doc, '数据集：HotpotQA（500 条，bridge 404 + comparison 96）作为主实验；2WikiMultiHopQA 作为边界分析，包括 100 条分题型诊断集与 300 条扩展检查。基线：Zero-Shot、Standard CoT、CoT-SC(N=3)、CoT-SC+GERS 重排、MoDeGraph-style 图提示基线。MoDeGraph-style 基线是依据公开方法描述实现的非官方版本，并非官方代码复现。本文方法：GERS+自适应、GERS-SC(K=3)、GERS-CV（+双向交叉核验）、GERS-CV2（+置信度加权，点估计最高）。模型 Qwen3-8B，4 worker 并行，零失败。指标：EM、F1、CS。主对比报告 EM/F1 差值的配对 bootstrap 95% CI，并对 EM 对错进行 McNemar 检验。')
    add_para(doc, '评估公平性保障：为降低评估链路缺陷扭曲对比的风险，本文统一三项处理——(1) 修复 EM 双向子串匹配 bug（原使数值答案虚高）；(2) 所有方法采用统一简洁答案提取与归一化；(3) GERS 汇总强制答案类型回扣。这些控制降低但不能完全消除 prompt 格式差异对比较的影响。')

    add_heading(doc, '4.2 主实验结果', 2)
    add_table(doc, ['方法', 'EM', 'F1', 'CS'], [
        ['Zero-Shot', '0.276', '0.389', '—'],
        ['Standard CoT', '0.264', '0.368', '—'],
        ['CoT-SC (N=3)', '0.262', '0.373', '—'],
        ['CoT-SC+GERS 重排', '0.264', '0.372', '—'],
        ['MoDeGraph-style prompt', '0.252', '0.366', '—'],
        ['GERS+自适应', '0.284', '0.395', '0.996'],
        ['GERS-SC (K=3)', '0.282', '0.398', '0.662'],
        ['GERS-CV（+双向交叉核验）', '0.298', '0.409', '0.782'],
        ['GERS-CV2（+置信度加权）', '0.302', '0.413', '0.777'],
    ], caption='表 1  主实验结果（HotpotQA, n=500）')
    add_para(doc, 'GERS-CV2 在本次运行中取得点估计最高 EM=0.302、F1=0.413，相比 CoT-SC EM +4pt、F1 +4pt，且 EM 对错的 McNemar 检验达到配对显著。它也显著高于本文非官方 MoDeGraph-style 图提示基线（F1 0.366；McNemar p=0.010，配对 F1 CI [+0.012,+0.083]），说明该 prompt-only 图实现是当前设置下的较弱基线，但不能外推到官方 MoDeGraph 或所有图推理系统。双向交叉核验（0.284→0.298）对应 +1.4pt EM 变化，是主要观察到的增益来源。HotpotQA 上 Zero-Shot（0.276）已接近 GERS，因上下文直接含答案证据压缩了结构化方法优势空间。')

    add_table(doc, ['对比', 'EM差值', 'EM 95% CI', 'F1差值', 'F1 95% CI', 'McNemar p'], [
        ['GERS-CV2 vs CoT-SC', '+0.040', '[+0.006,+0.074]', '+0.041', '[+0.006,+0.075]', '0.029'],
        ['GERS-CV2 vs StdCoT', '+0.038', '[+0.004,+0.072]', '+0.045', '[+0.010,+0.081]', '0.040'],
        ['GERS-CV2 vs MoDeGraph', '+0.050', '[+0.014,+0.086]', '+0.047', '[+0.012,+0.083]', '0.010'],
        ['GERS-SC vs CoT-SC', '+0.020', '[-0.012,+0.052]', '+0.025', '[-0.009,+0.060]', '0.275'],
    ], caption='表 2  显著性检验（HotpotQA, n=500）')
    add_para(doc, 'GERS-CV2 在 EM 对错的配对 McNemar 检验下显著优于 CoT-SC（p=0.029），且 EM/F1 差值的配对 bootstrap 95% CI 均不跨零（EM [+0.006,+0.074]、F1 [+0.006,+0.075]），即增益幅度虽小但配对显著。对 StdCoT（p=0.040）与 MoDeGraph-style 基线（p=0.010）同样显著。未引入双向交叉核验的 GERS-SC（p=0.275，CI 跨零）不显著，与核心机制的贡献方向一致。')

    add_heading(doc, '4.3 Consistency Score 区分度的修复（核心证据）', 2)
    add_table(doc, ['CS计算方式', '答对CS', '答错CS', '区分度', 'stdev'], [
        ['旧CS（纯结构, GERS-SC）', '0.6592', '0.6628', '-0.0035', '0.098'],
        ['新CS（+双向交叉核验, GERS-CV）', '0.7888', '0.7042', '+0.0847', '0.340'],
        ['对照：纯结构分 S_struct', '0.9967', '1.0000', '-0.0033', '—'],
    ], caption='表 3  CS 对错区分度（HotpotQA, n=500）')
    add_para(doc, '旧 CS 基本失效（区分度 -0.0035，反向噪声，分布聚集）。双向交叉核验修复了信号方向与幅度（区分度 +0.0847，分布拉开），但该信号仍不是正确性验证器。以 EM 对错为标签，旧结构 CS 的 AUROC 为 0.498，GERS-CV 为 0.581，GERS-CV2 为 0.589。GERS-CV2 中 250 条样本得到 CS=1.0，但其中 156 条仍为 EM 错误，说明内部自洽并不保证事实正确。')

    add_heading(doc, '4.4 分题型分析与适用边界', 2)
    add_table(doc, ['方法', 'comparison', 'bridge_comp', 'compositional', 'inference'], [
        ['Standard CoT', '0.815', '0.905', '0.284', '0.361'],
        ['Zero-Shot', '0.800', '0.587', '0.366', '0.330'],
        ['CoT-SC', '0.720', '0.864', '0.268', '0.343'],
        ['GERS-CV2', '0.775', '0.524', '0.297', '0.288'],
    ], caption='表 4  2WikiMultiHopQA 分题型 F1（n=100）')
    add_para(doc, 'comparison（纯对比题）：GERS-CV2 F1=0.775 与 Standard CoT 接近，DAG 拓扑分解有效拆解对比双方。bridge_comparison（桥接对比题）：GERS 短板（0.524 vs 0.905），子问题错误传播导致。双向交叉核验将 bridge_comparison F1 从 0.476（无CV）提升至 0.524，部分缓解但未根除，剩余差距源于真实推理错误传播。')
    add_para(doc, '证据约束验证（诊断实验）：前100条样本上，grounded 变体要求反向子答案引用上下文证据。HotpotQA 上 F1 基本持平/微升（GERS-CV2 0.451 vs grounded 0.455），CS 从 0.781 降至 0.730；2Wiki 上整体 F1 基本持平（0.469 vs 0.463），并在同一诊断重跑内将 bridge_comparison F1 从 0.524 提升到 0.571，但其他题型有抵消。由于这是单独诊断重跑，其分题型数值只应与同表 GERS-CV2 对比，不应和主 2Wiki n=100 分题型表混用。说明 evidence grounding 能抑制虚高 CS 并局部缓解桥接错误，但尚不足以形成稳定整体增益。')

    add_heading(doc, '4.5 计算成本分析', 2)
    add_table(doc, ['方法', '估计LLM调用/题', '单题延迟', 'EM', 'F1'], [
        ['Zero-Shot', '1', '0.7s', '0.276', '0.389'],
        ['Standard CoT', '1', '3.0s', '0.264', '0.368'],
        ['CoT-SC (N=3)', '3', '8.3s', '0.262', '0.373'],
        ['MoDeGraph-style', '3', '8.7s', '0.252', '0.366'],
        ['GERS+自适应', '1--4', '4.8s', '0.284', '0.395'],
        ['GERS-SC (K=3)', '~12--15', '16.2s', '0.282', '0.398'],
        ['GERS-CV2（最高）', '~6--8', '6.6s', '0.302', '0.413'],
    ], caption='表 5  计算成本对比（HotpotQA, n=500）')
    add_para(doc, 'GERS-CV2 单路版约需 6--8 次 LLM 调用，多于 CoT-SC 的 3 次采样，但在并行 API 设置下 wall-clock 延迟仍为 6.6s，且取得点估计最高 EM=0.302、F1=0.413。MoDeGraph-style 图提示基线同样约 3 次调用，但在本文实现中 F1 较低且耗时更高。当前测试的 GERS-SC（约 12--15 次调用，16.2s）在旧结构分选优下成本更高却无性能增益。')

    add_heading(doc, '4.6 消融实验', 2)
    add_table(doc, ['配置', 'EM', 'F1', '说明'], [
        ['GERS+自适应', '0.284', '0.395', '基线：DAG执行+自适应路由'],
        ['GERS-SC（+K=3）', '0.282', '0.398', '多路选优，旧CS无区分'],
        ['GERS-CV（+双向交叉核验）', '0.298', '0.409', '新CS有区分度'],
        ['GERS-CV2（+置信度加权）', '0.302', '0.413', '点估计最高'],
        ['w/o 图执行（=CoT-SC+GERS重排）', '0.264', '0.372', '无DAG执行'],
    ], caption='表 6  消融实验（HotpotQA, n=500）')
    add_para(doc, '双向交叉核验是主要观察到的增益来源（+1.4pt EM）。旧结构分图级自一致性无额外增益（p=1.000，针对旧分数 selector 的诚实负面发现）。置信度加权贡献有限（+0.4pt 不显著）。图执行是基础（移除后退化为 0.264）。')

    add_heading(doc, '4.7 案例分析', 2)
    add_para(doc, '案例 A（自洽且正确）：问题 Were Scott Derrickson and Ed Wood of the same nationality? 正向分解为两个导演国籍子问题，均得 American，最终答案 yes；反向验证以 yes+上下文为锚，独立重答也得到 American/American，crossval=1.0。')
    add_para(doc, '案例 B（不自洽并捕获幻觉）：问题 Which director had the longest career, Alain Resnais or Scott Sidney? 正向链幻觉出生/死亡年份并推出 Alain Resnais，但上下文缺少这些年份；反向验证从上下文重答时返回“上下文未提供”，7 个子问题全部不一致，crossval 降至 0.0。')
    add_table(doc, ['案例', '证据线索', '正向答案节选', '反向答案节选', 'Match', '新CS'], [
        ['A：国籍', 'Derrickson: American; Wood: American', 'American / American', 'American / American', '是', '1.0'],
        ['B：职业长度', '上下文缺少声称生卒年', '1922, 2014, 92yr, ...', '上下文未提供', '否', '0.0'],
    ], caption='表 7  真实案例：证据线索与正反向子答案一致性')

    # 5 讨论
    add_heading(doc, '5 讨论与局限', 1)
    add_para(doc, '双向交叉核验的核心价值：将 CS 区分度从 -0.0035 修复到 +0.0847，并将 AUROC 从 0.498 提升到约 0.58-0.59，使校验从"图是否合法"升级为"推理内容是否自洽"。该分数的价值主要是诊断与候选选择，而非完整事实验证。')
    add_para(doc, '性能的诚实定位：在配对评价下 GERS-CV2 显著优于 CoT-SC、StdCoT 与 MoDeGraph-style 基线（McNemar p≤0.040；配对 bootstrap CI 不跨零），但增益幅度较小（+4pt F1）。幅度有限因 HotpotQA 上下文直接含答案，Zero-Shot 已接近 GERS。')
    add_para(doc, '自洽不等于正确：CS 诊断与验证驱动局部修复的负面实验说明，提高 crossval 分数可以让推理链更内部自洽，但不必然提高 EM/F1。CS=1.0 仍包含大量错误答案，repair 变体也表现为 CS 上升而 EM/F1 下降。证据约束检查部分缓解了这一问题：它降低了虚高 CS，并提升了 2Wiki bridge_comparison，但仍未带来整体 F1 提升。')
    add_para(doc, '图级自一致性的边界：当前测试的 K=3 多路 selector 使用旧结构分，在中等难度多跳上无额外增益（p=1.000）。本文实验中的最佳成本--效果点是单路 DAG 执行加双向交叉核验；使用新 cross-validated score 扩展多路选优仍是未来工作。')
    add_para(doc, '深度复合桥接的局限：bridge_comparison 上 GERS 错误传播落后于线性 CoT，是图结构分解在深度多跳的固有局限。')
    add_para(doc, '图提示基线的覆盖范围：本文已在同一上下文、答案提取与指标口径下加入非官方 MoDeGraph-style 图提示基线。GERS-CV2 在 HotpotQA 上更高，但这只能说明当前 prompt-only 图实现弱于 GERS，不能外推为击败官方 MoDeGraph 或所有图推理系统。完整 GoT 或 RwG 风格实现及匹配采样仍超出本文当前范围。')

    # 6 结论
    add_heading(doc, '6 结论', 1)
    add_para(doc, '本文提出 GERS，核心创新是子答案双向交叉核验：以最终答案+上下文为锚反向重答子问题，对比正反向一致性，将 Consistency Score 从纯图论结构指标（区分度 -0.0035，AUROC 0.498）升级为更有信息量的内容自洽信号（区分度 +0.0847，AUROC 约 0.58-0.59）。该机制由显式 DAG 分解自然支持。点估计最高配置 GERS-CV2 在 HotpotQA 500 条上取得 EM=0.302、F1=0.413，F1 相比 CoT-SC 高 4pt，且在配对 McNemar 检验（p=0.029）与配对 bootstrap CI（不跨零）下均显著，并对 MoDeGraph-style 基线显著（p=0.010）。本文诚实呈现方法适用边界，并通过三项工程改进量化了"格式诱导错误"与"真实推理差距"的边界。总体而言，GERS 更应被视为可解释图推理诊断的一步，而非已经解决事实正确性验证。')
    add_para(doc, '未来工作：(1) 改进证据约束验证，引入更可靠的 span 抽取或外部检索信号；(2) 针对深度复合桥接题设计更保守的局部重生成机制，避免覆盖原本正确的中间答案；(3) 在更大规模模型与更多数据集上验证双向交叉核验的泛化性。')

    # 参考文献
    add_heading(doc, '参考文献', 1)
    refs = [
        '[1] Wei, J., et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models. In NeurIPS, 2022.',
        '[2] Besta, M., et al. Graph of Thoughts: Solving Elaborate Problems with Large Language Models. In AAAI, 2024.',
        '[3] Han, H., et al. Reasoning with Graphs: Structuring Implicit Knowledge to Enhance LLMs Reasoning. In Findings of ACL, 2025.',
        '[4] Oruche, R., et al. Disentangling Complex Questions in LLMs via Multi-Hop Dependency Graphs. In CIKM, 2025.',
        '[5] Wang, X., et al. Self-Consistency Improves Chain of Thought Reasoning in Language Models. In ICLR, 2023.',
        '[6] Yao, S., et al. Tree of Thoughts: Deliberate Problem Solving with Large Language Models. In NeurIPS, 2023.',
        '[7] Lightman, H., et al. Let\'s Verify Step by Step. In ICLR, 2024.',
        '[8] Madaan, A., et al. Self-Refine: Iterative Refinement with Self-Feedback. In NeurIPS, 2023.',
        '[9] Yang, Z., et al. HotpotQA: A Dataset for Diverse, Explainable Multi-hop QA. In EMNLP, 2018.',
        '[10] Ho, X., et al. Constructing A Multi-hop QA Dataset for Comprehensive Evaluation of Reasoning Steps. In COLING, 2020.',
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(r)
        set_font(run, size=9.5)

    out = Path('docs/paper.docx')
    doc.save(out)
    print(f'Word 版已生成: {out} ({out.stat().st_size//1024} KB)')


if __name__ == '__main__':
    main()
